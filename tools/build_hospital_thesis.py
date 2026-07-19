from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUT = PROJECT_ROOT / "Hospital_ICU_Transfer_DSS_Thesis.docx"


def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold=False, fill=None):
    if fill:
        set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    set_font(run, bold=bold)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color="666666")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True, color="2E74B5" if level < 3 else "1F4D78")
    return p


def add_para(doc: Document, text: str, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_font(run)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, bold=True, fill="F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


SECTIONS = [
    (
        "Abstract",
        [
            "Emergency inter-hospital transfer is a time-critical decision problem in which a sending hospital must identify a receiving hospital that can accept the patient, provide the required intensive care capability, and be reached safely by ambulance. In congested urban areas such as Colombo, the problem is not solved by choosing the geographically nearest hospital. A technically stronger solution must combine bed capacity, clinical urgency, hospital specialty, traffic conditions, road safety, ambulance location, and operational workflow status. This thesis presents the design and implementation of an ICU Capacity-Aware Emergency Transfer Decision Support System developed as a final year Computer Science project. The system is built as a full-stack web application with a FastAPI backend, SQLAlchemy data model, React and TypeScript frontend, OpenStreetMap road graph routing, machine-learning assisted traffic estimation, urgency scoring, ambulance dispatch, and administrative dashboards.",
            "The project implements a practical research prototype rather than a purely conceptual model. The backend exposes hospital, prediction, route, transfer, ambulance, traffic, and administration APIs. The frontend provides transfer planning, route visualization, bed management, request acceptance, mission control, live fleet display, and scenario simulation. The routing engine compares a shortest-time baseline with a proposed urgency-aware multi-objective route that balances estimated travel time, route risk, and distance. The local OSM graph allows road-following paths, node sequences, polylines, step instructions, and road-risk features such as traffic signals, intersections, high-risk segments, and local-road exposure. The dispatch model ranks available ambulances using pickup ETA, pickup risk, base coverage impact, same-base preference, stale-location penalty, and off-base penalty. Capacity forecasting and scenario analytics extend the system from a single transfer tool into a network-level operational decision support platform.",
            "The study follows an engineering methodology: problem analysis, requirement specification, system design, data modeling, algorithm design, implementation, testing, and evaluation. Because real clinical and ambulance traffic datasets are difficult to obtain, the project uses seeded Colombo hospital data, synthetic patient transfer data, cached traffic-shaped features, deterministic fallbacks, and transparent scoring rules. This approach keeps the prototype explainable and reproducible while still allowing machine learning artifacts to be integrated where data exists. The final result demonstrates how a hospital transfer DSS can support safer and more transparent emergency routing by combining deterministic filtering, predictive estimation, graph search, and human-readable explanations. The main contribution is an end-to-end prototype that links clinical urgency, capacity status, road-network intelligence, ambulance assignment, and hospital workflow into one coherent application.",
        ],
    ),
    (
        "Chapter 1: Introduction",
        [
            "Hospitals often operate as part of a wider emergency-care network rather than as isolated institutions. When one hospital lacks an available ICU bed, the patient must be transferred to another facility that can provide appropriate care. This decision is stressful because it is made under time pressure and with incomplete information. A dispatcher or hospital administrator must consider whether the destination has a suitable ICU type, whether ventilator support is available, whether the hospital can support the patient's condition category, whether an ambulance is available, and whether the route is acceptable for the urgency level. In traditional workflows these checks may be performed through phone calls, manual notes, and personal experience. That approach can work in simple cases, but it becomes unreliable when several hospitals, ambulances, and urgent requests compete for limited resources.",
            "The project described in this thesis addresses that operational gap by building a digital decision-support system for ICU-aware emergency transfers in Colombo. The selected domain is intentionally narrow: emergency inter-hospital transfers that require intensive care. Narrowing the domain made it possible to model the full decision chain in detail, from hospital capacity to ambulance mission completion. The application allows a user to choose an origin hospital, enter patient condition information, specify the required ICU type, receive ranked destination recommendations, request a transfer, wait for destination acceptance, trigger automated ambulance dispatch, and monitor the mission until completion. The design also includes administrative tools for updating ICU beds, viewing transfer timelines, forecasting capacity pressure, and running demand simulations.",
            "The central research problem is how to recommend a receiving hospital and ambulance route by combining clinical, operational, and geographic factors. A nearest-hospital strategy may choose a facility that lacks the required ICU type. A shortest-time route may pass through road segments with higher operational risk. A purely capacity-based strategy may overload the ambulance fleet or ignore congestion. Therefore, this project treats transfer planning as a multi-objective decision problem. The system filters candidate hospitals by suitability and availability, predicts urgency from patient condition features, estimates route travel time and risk, and ranks options through urgency-sensitive weights. Critical cases give greater importance to time, while moderate cases can give more weight to safer routing and capacity resilience.",
            "The developed prototype is not a replacement for medical judgement. It is a decision-support system that organizes information, calculates transparent scores, and presents explanations that a human operator can accept or reject. This distinction is important for both ethical and technical reasons. Clinical decisions require accountability, contextual knowledge, and professional authority. The system is designed to improve situational awareness and reduce manual coordination delays, not to remove human responsibility. The same design principle appears in the user interface: recommendation cards show reasons such as matching ICU capability, available beds, and urgency-adjusted route cost; route panels expose travel time, risk score, source, model used, and step instructions; transfer detail panels preserve event histories and patient handover data.",
            "The project uses a modern web architecture. The backend is written in Python with FastAPI, Pydantic schemas, SQLAlchemy models, and service classes that separate domain logic from route handlers. The database-backed MVP uses SQLite for local development and is designed to support PostgreSQL with PostGIS in production. The frontend is written in React, TypeScript, Vite, React Leaflet, MapLibre GL, and Lucide icons. The map view displays hospitals, ambulances, route polylines, and driver navigation modes. The machine learning layer uses scikit-learn pipelines and regressors for urgency and traffic experiments, while maintaining deterministic fallbacks for reproducibility. The routing layer uses NetworkX over a local OpenStreetMap GraphML file for Colombo.",
            "This thesis is organized as follows. Chapter 2 explains the background and related concepts behind hospital transfer systems, ICU capacity management, triage scoring, graph routing, route-risk analysis, and decision-support design. Chapter 3 defines the requirements and scope. Chapter 4 describes methodology. Chapter 5 presents the system architecture. Chapter 6 explains the data model and database design. Chapter 7 details the algorithms, including urgency scoring, transfer ranking, route optimization, ambulance dispatch, capacity forecasting, and simulation analytics. Chapter 8 covers implementation at module level. Chapter 9 discusses testing and evaluation. Chapter 10 summarizes results, limitations, future work, and conclusion. Appendices document API contracts, startup instructions, and implementation artifacts.",
        ],
    ),
    (
        "Chapter 2: Background and Related Work",
        [
            "Emergency medical transfer is a coordination problem that sits between clinical care, logistics, and information systems. A patient may be stabilized at one hospital but require a specialist service, ICU bed, ventilator, pediatric unit, maternity unit, neuro support, or trauma capability that is not currently available. The sending hospital must identify a receiving hospital, communicate patient details, obtain acceptance, dispatch transport, and ensure that the destination can prepare the bed. Each delay can affect patient outcome, especially in respiratory, cardiac, trauma, sepsis, and neuro cases. Information technology can reduce delay by making capacity, route, and ambulance information visible in one workflow.",
            "Capacity management is a major part of this problem. ICU beds are not simply interchangeable resources. A bed can be available, occupied, reserved, under cleaning, under maintenance, or assigned to an incoming transfer. A hospital may have general ICU beds but not pediatric or maternity ICU beds. Some hospitals may support ventilators while others may have limited respiratory capability. Capacity also changes over time because patients are discharged, beds are cleaned, transfers arrive, and emergency cases consume resources. For this reason, the project models both hospital-level bed counts and individual ICU bed lifecycle records. This micro-level modeling supports more accurate dashboards and audit trails than a single static available-bed number.",
            "Triage scoring is another key component. A transfer system must treat a critical ventilated respiratory patient differently from a stable moderate-priority case. In the implemented prototype, urgency is predicted through a rule-based scoring service. The score starts from a baseline and increases for high-risk condition categories, critical or low oxygen saturation, shock or unstable blood pressure, reduced consciousness or unconsciousness, and ventilator requirement. The rule-based approach is intentionally explainable and suitable for an MVP with limited real clinical data. The project also includes an ML training pipeline that compares Logistic Regression, Random Forest, and Gradient Boosting on synthetic transfer scenarios, allowing the same interface to support a trained urgency classifier later.",
            "Routing in emergency systems is also more complex than shortest path. A shortest route may include narrow local roads, many junctions, heavy signals, or structurally risky segments. A fastest route may be affected by congestion variability. The implemented system uses an OpenStreetMap road graph and NetworkX path search. Edge travel time is calculated from segment length, speed, and congestion ratio. Risk is calculated from graph features such as road class, intersections, traffic signals, and high-risk segment exposure. The route engine supports both a shortest-time strategy and an ML traffic risk-aware strategy, allowing the final evaluation to compare a baseline against the proposed multi-objective approach.",
            "Machine learning is used carefully. The project does not depend on a large black-box model for the entire decision. Instead, it follows a hybrid architecture: deterministic rules filter hospitals by ICU type, beds, condition support, and ventilator support; ML regressors estimate congestion ratio and duration when trained artifacts exist; OSM features provide route structure and risk; and a transparent multi-objective formula produces route cost. This hybrid approach is academically defensible because it uses machine learning where historical observations can support prediction while keeping safety-critical filtering explicit. The traffic model service also includes fallbacks based on time of day and road features, so the system remains usable without a Google API key or trained model files.",
            "Decision-support systems must present reasons, not only answers. If a ranked recommendation appears without explanation, a hospital administrator may not trust it. The implemented application therefore returns reasons and evidence at multiple levels. An urgency prediction includes the clinical factors that increased the score. A hospital recommendation includes matching ICU capability, available bed status, and selected route strategy. A route option includes distance, ETA, risk score, route source, model used, route steps, risk features, and risk factors. A dispatch decision includes score components and candidate rankings. This approach aligns with explainable decision-support principles and makes the prototype easier to evaluate in a final-year academic context.",
        ],
    ),
    (
        "Chapter 3: Requirements and Scope",
        [
            "The functional requirements were derived from the real workflow of a hospital transfer request. First, the system must maintain a dataset of Colombo hospitals with coordinates, ICU types, total beds, occupied beds, supported condition categories, and ventilator support. Second, it must accept patient condition inputs such as condition type, oxygen saturation band, blood pressure band, consciousness level, ventilator requirement, and required ICU type. Third, it must calculate urgency class and urgency score. Fourth, it must filter candidate receiving hospitals using ICU type, available beds, condition support, and ventilator compatibility. Fifth, it must rank candidate hospitals using capacity and route cost. Sixth, it must show route options and map-ready geometry.",
            "The transfer workflow adds another set of requirements. A hospital administrator must create a transfer request with patient demographic details, clinical notes, vitals, medications, allergies, infection risk, isolation requirement, and handover notes. The destination hospital must be able to accept or reject the request. On acceptance, the system must reserve a bed and automatically assign an ambulance when possible. The ambulance crew must be able to view its active mission, update location, start pickup, arrive at pickup, start destination travel, and complete delivery. Completion must update transfer status, ambulance availability, and assigned ICU bed occupancy. These requirements make the prototype operational rather than only advisory.",
            "Administrative requirements include a dashboard that shows hospitals, ambulances, transfers, beds, capacity forecast, simulation results, and alerts. Hospital admins should see information relevant to their hospital, while a super admin should see the whole network. Bed management must support status updates, patient assignment, patient removal, lifecycle events, and synchronization with hospital bed counts. The system also needs audit logging and transfer event timelines so that decisions can be reviewed after the fact. These features were implemented because hospital operations require traceability, especially when a transfer is rejected, accepted, assigned, or completed.",
            "Routing requirements focus on both technical and user-facing outcomes. The route engine must calculate a shortest-time baseline and a proposed risk-aware route. The backend must return estimated minutes, distance, risk score, total cost, congestion ratio, model used, route source, route node IDs, route steps, risk features, risk factors, explanation, and polyline coordinates. If the local OSM graph is available, it should produce road-following geometry. If the graph fails or is disabled, the system must fall back gracefully to cached traffic-model estimates or direct coordinate polylines. Fail-open behavior is important because a clinical API must not stop working only because one routing source is unavailable.",
            "Non-functional requirements include usability, responsiveness, explainability, maintainability, and local reproducibility. The application must start through provided Windows scripts and open a local frontend. The backend must initialize the database and warm up routing and traffic models during startup. The frontend must handle backend offline states and reconnect automatically. The service layer should be modular so that urgency, routing, transfer recommendation, dispatch, traffic prediction, fleet simulation, capacity forecasting, and simulation analytics can be tested independently. The project also requires simple deployment support through Docker Compose and environment-variable configuration for database URLs, Google API keys, OSM routing toggles, and fleet simulation toggles.",
            "The scope is intentionally limited to a research prototype. It does not integrate with real hospital electronic medical records, real ambulance GPS feeds, national health registries, or live production traffic services by default. It uses seeded hospital data and synthetic or cached data for experiments. It is designed to demonstrate architecture, algorithms, and workflow feasibility rather than to operate as a certified medical device. This limitation is important because medical software requires clinical validation, cybersecurity hardening, privacy review, and regulatory approval before real deployment. The final-year project contribution is therefore a technically complete prototype and research platform, not a production clinical system.",
        ],
    ),
    (
        "Chapter 4: Methodology",
        [
            "The methodology followed an iterative software engineering approach. The first phase defined the research problem and decomposed it into capacity, triage, routing, transfer workflow, ambulance dispatch, and visualization modules. The second phase created a minimal working backend with hospital data, urgency prediction, and route recommendation. The third phase added a React frontend with map visualization and transfer planning. The fourth phase expanded the prototype into a database-backed operational system with users, ambulances, ICU beds, patient records, transfer events, and bed lifecycle events. The final phase integrated OSM graph routing, traffic model support, dispatch scoring, capacity forecasting, simulation analytics, and driver navigation.",
            "The project uses a hybrid design methodology because no single technique solves the whole problem. Deterministic rules are used where the domain requires hard constraints: a hospital without the requested ICU type should not be recommended, a hospital without available beds should be excluded, and a ventilated patient should not be sent to a hospital without ventilator support. Predictive models are used where uncertainty exists: travel time and congestion vary by time of day and route. Graph algorithms are used for road routing because road networks are naturally represented as nodes and edges. Scoring formulas combine these signals into ranked choices that remain inspectable.",
            "Data preparation was performed at several levels. The `data/seed_hospitals.json` file defines nine Colombo hospitals with coordinates, ICU types, bed counts, capability flags, and ventilator support. The seed database expands this into hospital records, admin users, ambulance crew users, ambulances, ICU beds, and demo patient records. The ML folder contains scripts to generate synthetic transfer data for urgency classification and to train traffic models using traffic-ready features. The traffic model expects features such as hospital pair, time, weekend flag, morning and evening peak flags, sine/cosine encodings, static duration, distance, OSM intersection count, signal count, road mix, and traffic interval ratios.",
            "Implementation followed a layered backend structure. FastAPI route handlers receive requests, validate schemas, check user scope, and call service classes. Services perform domain logic. Models define database tables. Schemas define typed API contracts. This organization avoids placing all logic inside endpoints and makes the code easier to explain. For example, `TransferService` handles recommendation logic, `RoutingService` handles route comparison, `OSMGraphRoutingService` handles graph search, `DispatchService` ranks ambulances, `CapacityForecastService` forecasts bed pressure, and `SimulationAnalyticsService` calculates scenario impact. This separation is one of the project’s main engineering strengths.",
            "The frontend methodology emphasized workflow completeness. Rather than building only a static map, the UI models multiple roles and operational views. A super admin can observe the network, hospital admins can manage transfers and beds, and ambulance crew can view active missions. The application stores state for hospitals, users, recommendations, routes, dashboard data, capacity forecasts, simulation data, mission data, backend online status, selected transfer detail, transfer events, ICU beds, and map command views. React hooks handle data fetching, loading states, errors, modal confirmations, and automatic updates. Leaflet supports the standard operations map, while MapLibre GL provides a 3D driver-style navigation view.",
            "Testing and evaluation were planned around functional correctness and decision quality. Functional tests involve endpoint calls, route optimization, transfer creation, transfer acceptance, automatic ambulance assignment, mission actions, bed updates, and dashboard refreshes. Decision quality is evaluated through whether recommendations satisfy ICU type, bed availability, condition capability, ventilator support, reasonable travel time, and lower urgency-adjusted route cost. The route evaluation compares shortest-time and ML traffic risk-aware strategies. Capacity and simulation outputs are checked against expected pressure levels and recommended actions. Because the dataset is controlled, the project prioritizes transparent behavior and reproducible scenarios over claims of clinical outcome improvement.",
        ],
    ),
    (
        "Chapter 5: System Architecture",
        [
            "The architecture is a full-stack client-server system. The frontend runs as a Vite React application on the local development port, while the backend runs as a FastAPI application on port 8001 or 8000. The frontend communicates with backend endpoints using a helper that tries configured API base candidates and stores the active base once a request succeeds. This makes local startup more tolerant of port differences. The backend exposes routes under `/api/hospitals`, `/api/admin`, `/api/ambulance`, `/api/predictions`, `/api/routes`, `/api/traffic`, and `/api/transfers`. It also provides `/health` for readiness checks and redirects `/` to the generated API documentation.",
            "At startup, the backend initializes the database, starts model warmup in a background thread, and starts the fleet simulation service. Model warmup preloads traffic features, congestion models, duration models, and OSM graph routing. This design reduces first-click latency in the user interface. Fleet simulation runs in a daemon thread and moves available off-base ambulances back to their base hospitals while keeping them assignable. The simulation is deliberately fail-safe: exceptions inside the simulation loop are swallowed so that the clinical API continues running. Environment variables can disable OSM routing or fleet simulation when connecting to real routing services or live GPS feeds.",
            "The backend is divided into route, service, model, schema, database, authentication, and seed modules. Route files handle HTTP concerns. Service files contain the core algorithms. The database module creates the SQLAlchemy engine and session. The authentication module implements user resolution and role checks. The seed module populates demo hospitals, users, ambulances, beds, and baseline data. This structure supports a clean explanation in the thesis because each module maps to a responsibility. It also allows future replacement: SQLite can be replaced by PostgreSQL, the rule-based urgency service can be replaced by a trained model, and fleet simulation can be replaced by live ambulance telemetry.",
            "The frontend architecture is a single main application component with typed domain models and a lazily loaded 3D navigation component. Although the main `App.tsx` file is large, it implements a coherent console with multiple command views: live map, fleet, capacity, beds, transfers, analytics, and alerts. The app uses React Leaflet to render hospitals and ambulances as map markers, route polylines as highlighted paths, and popups with hospital or ambulance information. When an ambulance crew user is active and a mission route exists, the UI can switch into driver mode and render a MapLibre GL scene with pitch, bearing, 3D buildings, ambulance marker, pickup/dropoff markers, and route following.",
            "The data flow starts with hospital and user loading. A user selects an origin hospital, patient condition, and required ICU type. The frontend posts to the recommendation endpoint. The backend predicts urgency, filters candidate hospitals, compares route strategies for each candidate, calculates scores, and returns ranked recommendations. If the user requests a transfer, the admin endpoint persists a transfer request with patient and handover data. The destination admin accepts the request. The backend reserves an ICU bed, calls the dispatch service, assigns an ambulance, stores pickup and destination route payloads, writes audit logs, and emits transfer events. The ambulance crew then progresses through mission actions until delivery.",
            "A key architectural quality is fallback behavior. The route service prefers the local OSM graph when available. If the graph is unavailable, it can use live Google Routes when configured. If live routes are unavailable, it uses cached Google-shaped features and trained or fallback traffic models. If no detailed polyline exists, it returns direct coordinate fallback geometry. This layered design is important for a final-year prototype because it allows the application to run reliably on a local machine while still demonstrating how real traffic and routing providers can be integrated in production. The user interface exposes route source and model used so the operator understands which data path produced the result.",
        ],
    ),
    (
        "Chapter 6: Database and Data Model",
        [
            "The database design evolved from the original blueprint into a practical operational schema. The current SQLAlchemy models include hospitals, users, ambulances, ICU beds, patient records, transfer requests, transfer events, bed lifecycle events, and audit logs. The original documentation also describes production-oriented tables such as ICU units, ICU bed status, route options, route risk features, and PostGIS geometry columns. The current MVP uses SQLite for local development, but the schema is compatible with a future PostgreSQL and PostGIS deployment through the `DATABASE_URL` environment variable. This design supports quick demonstration while preserving a migration path.",
            "The hospital model stores stable hospital identifiers, names, coordinates, ICU type, total beds, occupied beds, capability flags, ventilator support, phone, address, and update time. The seeded dataset includes National Hospital Sri Lanka, Colombo South Teaching Hospital, Lady Ridgeway Hospital, Sri Jayewardenepura General Hospital, De Soysa Hospital for Women, Nawaloka Hospital, Asiri Central Hospital, Lanka Hospital, and Durdans Hospital. Each hospital has ICU metadata and support flags for trauma, cardiac, neuro, pediatric, maternity, and ventilator-related decisions. The transfer recommendation service uses these fields directly when it filters candidates.",
            "The user model supports role-based workflow simulation. Users include a super admin, hospital admins, and ambulance crew users. A hospital admin is scoped to a hospital, while an ambulance crew user is scoped to an ambulance. This allows the frontend to simulate different operational roles without a complex authentication provider. Role checks ensure that a user can only view or act on transfers within their scope unless they are a super admin. This design is important because real hospital systems require separation between sending hospitals, receiving hospitals, dispatch administrators, and ambulance crews.",
            "The ambulance model stores call sign, base hospital, status, current latitude and longitude, crew contact, and update time. Status values include available, assigned, en route, transporting, returning, repositioning, and offline. The dispatch service queries available ambulances and ranks them for a transfer. The fleet simulation service keeps available ambulances parked at their base or moving home along an OSM route. The ambulance route payload stores both pickup and destination legs because crew navigation changes after the patient is onboard. This micro-level route separation is more realistic than a single straight-line mission route.",
            "The ICU bed model provides bed-level state. Each bed has an identifier, hospital, bed number, ICU type, ward, status, FHIR location ID, operational status, status reason, and update time. Status values include available, occupied, transfer assigned, reserved, cleaning, and maintenance. A bed can be linked to a patient record. The admin update endpoint can assign a patient to a bed, clear a patient, update status reason, and record lifecycle events. Hospital occupied-bed counts are synchronized from bed statuses, reducing inconsistency between summary and bed-level data. This is one of the project’s more detailed implementation improvements compared with the initial blueprint.",
            "The patient record and transfer request models store the clinical packet needed for handover. Patient records include patient number, name, identifier, date of birth, age, sex, blood type, condition, diagnosis, vitals, medications, allergies, infection risk, isolation requirement, emergency contact, next of kin, address, notes, and admission/update timestamps. Transfer requests copy relevant patient details, requested ICU type, urgency class, urgency score, ventilator requirement, handover JSON, ambulance assignment, assigned bed, route payload, pickup/dropoff coordinates, status, and timestamps. The model therefore connects clinical inputs, operational routing, and final bed occupancy into one traceable lifecycle.",
            "Event tables support accountability. Transfer events record status changes such as request creation, acceptance, ambulance assignment, mission progress, rejection, and completion. Bed lifecycle events record status transitions, reasons, patient IDs, and transfer IDs. Audit logs record user actions against entities with optional JSON details. These tables are not merely administrative extras; they are essential for explaining why a recommendation or workflow action occurred. In a hospital environment, decision histories matter for quality review, training, and incident analysis. The thesis therefore treats event logging as part of the decision-support architecture rather than as a secondary logging feature.",
        ],
    ),
    (
        "Chapter 7: Algorithms and Decision Logic",
        [
            "The urgency algorithm is implemented as a transparent rule-based score. It begins with a baseline score of 0.20. High-risk condition categories such as cardiac, trauma, respiratory, and sepsis add 0.15. Critical oxygen saturation adds 0.30, while low oxygen saturation adds 0.20. Shock adds 0.25 and unstable blood pressure adds 0.15. Unconsciousness adds 0.20 and reduced consciousness adds 0.10. Ventilator requirement adds 0.20. The score is capped at 1.00. Scores of 0.75 or above are classified as critical, scores of 0.50 or above are high, and lower scores are moderate. The service returns both the class and the explanation list.",
            "The hospital filtering algorithm applies hard constraints before scoring. The candidate cannot be the origin hospital. The normalized required ICU type must be present in the hospital's ICU types. Available beds must be greater than zero. The hospital must support the patient condition category through its capability flags. If the patient requires a ventilator, the hospital must have ventilator support. These checks prevent the scoring model from hiding unsafe choices behind a numerical ranking. For example, a hospital with excellent route time but no pediatric ICU should not appear for a pediatric ICU request. The filter stage therefore represents domain safety logic.",
            "The route comparison algorithm returns two options for an origin-destination pair: shortest time and ML traffic risk-aware. It first asks the traffic model service for distance, static duration, predicted duration, congestion ratio, risk score, and model source. It then asks the local OSM graph route service for road-following geometry. When the graph exists, the route includes node IDs, polyline, steps, distance, ETA, risk score, risk features, and risk factors. If the graph is not available, the route service falls back to Google Routes or cached traffic-shaped estimates. Both strategies are converted into a common `RouteOption` schema so the frontend can compare them uniformly.",
            "The OSM routing algorithm uses a directed NetworkX graph loaded from `data/osm/colombo_drive.graphml`. Each node has coordinates and each edge may contain length, speed, road name, road class, traffic signals, intersections, and risk. The service finds the nearest graph nodes to the origin and destination using coordinate distance adjusted by latitude scale. It then runs A* path search with a zero heuristic, effectively using weighted Dijkstra behavior while preserving the A* interface. For shortest-time routing, edge weight is travel seconds. For risk-aware routing, edge weight is travel seconds plus risk multiplied by distance units and an urgency-specific risk-seconds factor.",
            "The risk-aware edge weight is important. Critical cases use a lower risk penalty because speed is dominant; moderate cases use a higher risk penalty because the system can afford to avoid risky segments. The implemented risk-seconds factors are 2.0 for critical, 5.0 for high, and 9.0 for moderate. Travel seconds are calculated from edge length, speed in kilometers per hour, and congestion ratio. This creates an interpretable tradeoff between time and road risk. The output route is then summarized by distance-weighted average edge risk, high-risk kilometers, traffic signal count, intersection count, primary-road kilometers, residential-road kilometers, step count, and node count.",
            "The final route cost combines normalized time, risk score, and normalized distance. The weights are urgency-sensitive: critical uses time 0.75, risk 0.20, and distance 0.05; high uses time 0.60, risk 0.30, and distance 0.10; moderate uses time 0.45, risk 0.40, and distance 0.15. Time is normalized by 45 minutes and distance by 20 kilometers, each capped at 1.00. The formula produces a total cost in which lower is better. The transfer recommendation score is then calculated as one minus the best route total cost plus a small capacity score bonus based on available beds over total beds. Recommendations are sorted descending.",
            "The ambulance dispatch algorithm ranks available ambulances after a transfer is accepted. It first shortlists the eight nearest available ambulances to the origin hospital by Haversine distance. Each candidate is scored using OSM pickup route ETA when available, pickup risk, base coverage penalty, same-origin base bonus, stale-location penalty, and off-base penalty. Critical cases give lower urgency multipliers to ETA because every minute matters; moderate cases tolerate more caution. Coverage penalty is higher if assigning an ambulance leaves its base with no spare unit. A same-base ambulance receives a negative bonus because it is operationally natural for a hospital-based unit to serve its own request.",
            "The dispatch score components are stored in the transfer route payload. They include ETA cost, risk cost, coverage cost, base match bonus, stale location penalty, and off-base penalty. The selected ambulance is assigned, its status becomes assigned, and the transfer status becomes ambulance assigned. The payload also stores candidate rankings for the top candidates. This is valuable for transparency because an operator can inspect why a cross-base ambulance was selected or why a nearby ambulance was penalized due to coverage impact. The system therefore treats dispatch as an explainable optimization problem rather than a hidden nearest-vehicle lookup.",
            "The capacity forecasting algorithm estimates network pressure at 1, 3, 6, and 12 hour horizons. For each hospital it reads current bed statuses, counts available and occupied beds, counts active inbound transfers, estimates recent release rate from bed lifecycle events, and estimates cleaning recovery capacity from cleaning beds. Expected arrivals are based on inbound transfers, while expected releases combine discharge and cleaning recovery rates. Predicted available beds are clamped between zero and total beds. Pressure score is occupied beds over total beds. Pressure levels are stable, elevated, high, and critical, with recommended actions such as diverting non-critical transfers or expediting cleaning.",
            "The scenario simulation algorithm extends forecasting by injecting demand profiles. It supports baseline, evening surge, mass casualty, and respiratory wave scenarios. Each scenario has arrivals per hospital per six hours, critical share, ambulance factor, and description. The service scales arrivals by duration and intensity, calculates critical transfers, ambulances required, ambulance gap, total bed shortage, hospital-level impacts, pressure levels, and recommended actions. Arrival distribution uses hospital weights based on available capacity and total size, with different weights for mass casualty and respiratory scenarios. This gives administrators a read-only stress test that can be used for demonstration and planning.",
        ],
    ),
    (
        "Chapter 8: Implementation Details",
        [
            "The backend entry point defines a FastAPI application titled `ICU Transfer Decision Support API`. CORS is configured for local frontend origins on ports beginning with 517. Routers are included for hospitals, admin, ambulance, predictions, routes, traffic, and transfers. The startup function initializes the database, starts warmup models in a daemon thread, and starts the fleet simulation service. The warmup function touches traffic feature lookup, congestion model, duration model, and OSM graph warmup. This implementation detail matters because it shifts expensive loading away from the first user interaction and improves perceived responsiveness.",
            "The hospital API returns summary data such as ID, name, coordinates, ICU types, total beds, occupied beds, available beds, capability flags, and ventilator support. Filtering can be applied by ICU type, available-bed flag, and condition type. The prediction API exposes a single urgency endpoint that accepts condition fields and returns urgency class, score, and explanation. The routes API exposes route optimization between two hospitals and returns the two route strategies. The transfers API exposes recommendation logic for the older lightweight workflow, while the admin API implements the database-backed operational workflow used by the final prototype.",
            "The admin route file is the largest backend controller because it coordinates many operational actions. It includes helpers to convert models into schema summaries, parse JSON payloads, synchronize bed counts, record bed lifecycle events, reserve transfer beds, create patient records, serialize transfer summaries, add transfer events, audit actions, and check transfer scope. Dashboard endpoints return hospitals, ambulances, and transfers. Forecast and simulation endpoints call service classes and convert results into Pydantic response models. Streaming admin events support live-ish UI updates. Additional endpoints list users, list transfer events, update hospitals, list/update ICU beds, list/update ambulances, create transfers, accept transfers, reject transfers, and assign ambulances.",
            "The accept-transfer implementation is a critical workflow. It validates that the transfer exists and is pending destination acceptance. It checks the accepting user’s scope. It reserves an available ICU bed at the destination by changing the bed status to transfer assigned and creating a bed lifecycle event. It updates the transfer to accepted pending ambulance and records acceptance. It then calls the dispatch service. If dispatch succeeds, the transfer becomes ambulance assigned, the ambulance status becomes assigned, pickup and destination route payloads are stored, dispatch score details are attached, and transfer events are created. If no ambulance is available, the transfer remains accepted pending ambulance, making the situation visible instead of failing silently.",
            "The ambulance API supports crew operations. `GET /api/ambulance/mission` returns the crew’s ambulance, active transfer, and optional return route JSON. The mission action endpoint maps actions to transfer and ambulance statuses. `start-pickup` changes the transfer to ambulance en route to pickup and the ambulance to en route. `arrive-pickup` changes the transfer to patient onboard and the ambulance to transporting. `complete` changes the transfer to completed, returns the ambulance to available, marks the assigned bed occupied, creates or updates the patient record, and records lifecycle and transfer events. This endpoint closes the loop from recommendation to real resource consumption.",
            "The traffic model service is designed for multiple operating modes. It searches for model artifacts in `ml/artifacts` and an imported fallback folder. It loads a congestion ratio model, duration model, and feature dataset if available. It builds feature rows using exact hospital pair and time buckets when possible, or fallback geometry when not. Feature engineering includes time encodings, weekend flags, peak flags, traffic interval counts, jam ratios, slow-or-jam ratios, static ETA over distance, hospital pair code, and traffic band code. If models are missing, a time-of-day fallback congestion ratio is used: higher during morning and evening peaks, lower at night, and moderate during normal hours.",
            "The machine learning training scripts support reproducible experimentation. `train_urgency_model.py` loads synthetic transfer data, builds a scikit-learn pipeline with one-hot encoding for categorical features and passthrough boolean features, trains Logistic Regression, Random Forest, and Gradient Boosting classifiers, evaluates accuracy and classification reports, and saves the best model. `train_traffic_models.py` loads traffic-ready data, verifies required columns, checks available disk space, trains Random Forest, Extra Trees, and Gradient Boosting regressors for congestion ratio and duration seconds, selects the model with lowest mean absolute error, and writes artifacts plus a report. These scripts make the thesis more than a hand-coded system; they provide an experimental ML pathway.",
            "The frontend implements several interface modes. The transfer planner allows user selection, origin hospital selection, patient condition entry, ICU requirement selection, recommendation generation, route loading, and transfer request creation. The console includes overview, transfer, requests, and fleet tabs. The map command view can switch among live, fleet, capacity, beds, transfers, analytics, and alerts. Capacity panels display forecast horizon values. Simulation panels allow scenario, duration, and intensity selection. Bed panels allow ICU bed editing and patient assignment. Transfer detail panels show status, urgency, assigned bed, dispatch decision, patient packet, and event timeline. This breadth supports the thesis claim that the system implements an end-to-end workflow.",
            "The map implementation is central to the user experience. Standard mode uses React Leaflet with OpenStreetMap tiles, hospital circle markers, ambulance custom markers, route polylines, popups, and map fitting behavior. Ambulance markers are styled according to status and base. Driver mode uses the `ThreeDNavigationMap` component, which creates a MapLibre GL map with pitch, bearing, route shadow, route line, ambulance marker, pickup/dropoff/base markers, and recenter behavior. The component calculates bearings from route points and smoothly follows the ambulance location when following is enabled. This creates a visually clear difference between administrator overview and crew navigation.",
            "The startup scripts improve usability for demonstration. `START_APP.cmd` starts both backend and frontend, waits for health checks, and opens the local application. `START_OSM_APP.cmd` restarts with local OSM node routing enabled. `START_LEGACY_APP.cmd` restarts with previous Google or ML routing behavior. Environment variables control OSM graph routing and fleet simulation. The README warns not to run only the frontend because hospital dropdowns, recommendations, beds, ambulances, and dashboards require the backend. This practical packaging is important for a final-year presentation because examiners and supervisors should be able to run the prototype without learning every internal command.",
            "Styling and interaction design were implemented with a professional operational dashboard tone. The interface uses dense but readable panels, map-first layout, status chips, modal confirmations, loading indicators, backend offline banners, and role-specific command options. Error states are shown when the backend is unreachable or actions fail. The UI avoids making the transfer planner a marketing page; instead, the first screen is the usable command environment. This matches the domain: emergency transfer support should be direct, fast, and scannable. The visual design therefore supports the system’s purpose rather than distracting from it.",
        ],
    ),
    (
        "Chapter 9: Testing and Evaluation",
        [
            "Testing was organized around the core workflows. The first test is startup: running the provided application script should initialize the backend database, start the FastAPI server, start the Vite frontend, and open the local URL. The `/health` endpoint should return an OK status. The frontend should load hospitals, users, dashboard data, capacity forecast, and fleet information. If the backend is offline, the UI should show a reconnecting message instead of failing silently. This verifies that the system is usable as a local demonstrator.",
            "Endpoint testing verifies individual modules. `GET /api/traffic/status` should report whether a Google API key is enabled, whether congestion and duration models are available, whether feature data is available, and how many rows are loaded. `POST /api/predictions/urgency` should return expected classes for clinical examples. A respiratory patient with low oxygen saturation, unstable blood pressure, reduced consciousness, and ventilator requirement should receive a high or critical score with explanations. `POST /api/routes/optimize` should return both shortest-time and ML traffic risk-aware options with route source, model used, ETA, distance, risk, total cost, and polyline.",
            "Recommendation testing checks safety constraints. When an origin hospital, required ICU type, and patient condition are submitted, the system should exclude the origin hospital, hospitals without the required ICU type, hospitals without available beds, hospitals that do not support the condition category, and hospitals without ventilator support when ventilator is required. Returned recommendations should contain ranks, destination names, scores, available beds, estimated minutes, route risk scores, and reasons. The recommended order should change when urgency changes because route weights change. Critical cases should favor time more strongly, while moderate cases should tolerate safer but possibly slower routes.",
            "Workflow testing follows the demonstration guide. A user selects National Hospital Admin, runs a recommendation, and requests transfer to a recommended hospital. Another user, such as Durdans Hospital Admin, accepts the pending transfer. The backend reserves a destination ICU bed and automatically assigns an ambulance if available. The ambulance crew user then opens the mission view, sees pickup and dropoff route data, starts pickup, marks arrival at pickup, starts or continues to dropoff, and completes the mission. After completion, the transfer is completed, the ambulance becomes available, and the assigned ICU bed becomes occupied. Transfer events should document the complete timeline.",
            "Dispatch evaluation focuses on score transparency. For a selected transfer, the transfer detail panel should show the dispatch model, score, pickup ETA, pickup risk, coverage impact, and explanation reasons. The route payload should include candidate rankings so the selected ambulance can be compared with alternatives. A good dispatch decision is not always the nearest vehicle; it may consider whether assigning the nearest ambulance leaves another hospital uncovered. The evaluation therefore checks that the dispatch score components match the intended logic and that explanations are understandable to a human operator.",
            "Routing evaluation compares the baseline and proposed route strategies. The shortest-time route should minimize estimated travel seconds. The ML traffic risk-aware route should include risk penalties based on urgency and OSM graph features. The evaluation checks route distance, ETA, risk score, total cost, risk features, and route steps. It also checks fallback behavior by disabling OSM routing or running without a Google key. The expected result is that the system still returns valid routes with clear route-source labels. This reliability is important because a decision-support system should degrade gracefully.",
            "Capacity forecasting evaluation checks whether forecast points respond logically to current beds, inbound transfers, recent releases, and cleaning beds. A hospital with many occupied beds and inbound transfers should show higher pressure. A hospital with recent bed releases or cleaning recovery should show improved future availability. Pressure levels should move from stable to elevated, high, and critical according to thresholds. Recommended actions should match the worst forecast point. Simulation evaluation checks whether baseline, evening surge, mass casualty, and respiratory wave scenarios produce different transfers, critical share, ambulance requirements, bed shortages, and network actions.",
            "Frontend evaluation checks usability and visual correctness. Hospitals and ambulances should appear at their coordinates, route polylines should follow roads when OSM routing is available, popups should show relevant data, and driver mode should show route, markers, and recenter behavior. Loading spinners should appear during requests. Blocking modals should prevent accidental irreversible actions such as transfer acceptance or mission completion. Detail panels should not hide patient safety information. The UI should remain readable when data is missing, when no simulation has been run, or when there are no transfer events.",
            "The project’s evaluation is limited by dataset realism. The hospital data is seeded for demonstration. Patient urgency data is synthetic or rule-based. Traffic features may be cached or generated from available files. No real clinical outcome dataset was used, and no formal user study was conducted with hospital staff. Therefore, evaluation claims are framed around system correctness, algorithmic transparency, workflow completeness, and prototype feasibility. The thesis does not claim that the system reduces mortality or transfer time in practice; it demonstrates a technical approach that could be validated with real data in future work.",
        ],
    ),
    (
        "Chapter 10: Results, Limitations, and Future Work",
        [
            "The main result of the project is a working end-to-end prototype for ICU capacity-aware emergency transfer support. The system can load a Colombo hospital dataset, calculate urgency, filter receiving hospitals, compare route strategies, recommend destinations, create transfer requests, reserve beds, assign ambulances, guide ambulance missions, update bed occupancy, display dashboards, forecast capacity pressure, and simulate demand scenarios. This is a broader result than the original MVP, which focused mainly on choosing an origin, entering patient condition, ranking hospitals, and showing a route. The final implementation covers the entire operational transfer lifecycle.",
            "A second result is the integration of multiple computational methods into one workflow. The project combines rule-based triage, deterministic capability filtering, ML traffic estimation, OSM graph routing, multi-objective route cost, dispatch scoring, bed lifecycle tracking, and scenario simulation. Each method is used where it is appropriate. This avoids the common mistake of forcing machine learning into every part of the system. The final architecture shows that practical decision support often requires a hybrid of rules, models, graphs, and user interface design. This is a valuable lesson for a Computer Science final-year project.",
            "A third result is explainability. The system returns urgency explanations, recommendation reasons, route risk factors, route steps, dispatch candidate rankings, score components, transfer events, and bed lifecycle events. These outputs make the system inspectable. A hospital user can understand why a destination was recommended, why a route was selected, why an ambulance was assigned, and how a bed changed status. Explainability is not only useful for trust; it also makes testing easier because expected behavior can be compared against visible reasons and values.",
            "The project has several limitations. The first limitation is data quality. Seeded hospital data and synthetic patient data are sufficient for demonstration but not for clinical validation. The second limitation is real-time integration. The prototype can collect live traffic snapshots if a Google API key is configured, but it does not continuously ingest live ambulance GPS, hospital bed feeds, or electronic medical records. The third limitation is security. The role model is suitable for local demonstration, but production deployment would require strong authentication, authorization, encryption, audit retention, and privacy controls. The fourth limitation is regulatory readiness; the prototype is not certified medical software.",
            "There are also algorithmic limitations. The OSM graph route service currently generates a single path per strategy rather than a full set of Pareto-front alternatives. The risk score is engineered from structural road features and does not use historical ambulance incident labels. The traffic model depends on available cached or trained artifacts and may fall back to formula-based congestion estimates. The dispatch model uses transparent heuristic scoring rather than a formally optimized fleet model. These limitations are acceptable for a final-year prototype but identify clear research directions for future versions.",
            "Future work should begin with real data integration. Hospitals could expose bed status through standards such as HL7 FHIR Location and Encounter resources. Ambulances could provide GPS updates through a secure telemetry API. Traffic providers could stream live travel time estimates. Historical transfer records could support better urgency, ETA, risk, and dispatch models. With sufficient labels, route risk could be trained from actual delay, incident, or expert-review outcomes rather than only road structure. A production version should also store route geometries in PostGIS and use OSRM, Valhalla, GraphHopper, or pgRouting for faster route generation.",
            "Future user interface work should involve hospital stakeholders. Administrators, ICU coordinators, dispatchers, and ambulance crews should test the workflow and provide feedback on terminology, alert thresholds, screen density, and mission actions. Usability studies could measure time to create a transfer, time to identify a destination, error rate in patient handover entry, and perceived trust in recommendations. The simulation module could be extended into a planning dashboard for exercises, surge response, and daily capacity huddles. Exportable reports could support management review and academic evaluation.",
            "Future security and governance work is essential. The system handles patient identifiers, diagnoses, vitals, medications, allergies, infection risk, and emergency contacts. Production deployment would need strict privacy controls, role-based access, consent and retention policies, secure audit logs, encrypted transport, encrypted storage, backup and recovery, and integration governance. The prototype already contains audit logs and role scopes, but these are foundations rather than complete safeguards. A real deployment would also require clinical safety cases, hazard analysis, failover planning, and human override design.",
            "In conclusion, the ICU Capacity-Aware Emergency Transfer DSS demonstrates how software engineering, data modeling, graph algorithms, machine learning, and user interface design can be combined to address a realistic healthcare coordination problem. The project moves beyond a simple hospital finder by modeling patient urgency, ICU capacity, route risk, ambulance dispatch, mission status, bed lifecycle, and network pressure. Its main academic contribution is a transparent, modular, and runnable prototype that can be inspected, tested, and extended. With real data, stronger security, and clinical validation, the approach could support safer and more efficient inter-hospital emergency transfers.",
        ],
    ),
    (
        "Chapter 11: Security, Ethics, Deployment, and Project Management",
        [
            "A hospital transfer system handles information that is operationally urgent and personally sensitive. Even in a prototype, it is necessary to discuss privacy and security because the implemented data model includes patient names, identifiers, date of birth, age, sex, blood type, diagnosis, vitals, medications, allergies, infection risk, isolation requirement, emergency contact, and handover notes. In a real deployment these fields would be protected health information. The prototype stores them locally for demonstration, but the design points toward stronger controls: authenticated users, scoped roles, audit logs, transfer events, and bed lifecycle events. These foundations show awareness that healthcare software must control who can read or change data and must preserve evidence of critical actions.",
            "Role-based access is the first security boundary in the implemented system. A super admin can inspect the network, a hospital admin is associated with a hospital, and an ambulance crew user is associated with an ambulance. The backend checks whether a user can see or act on a transfer by comparing the transfer's origin hospital, destination hospital, and ambulance assignment with the user scope. This prevents a hospital user from casually managing unrelated transfers in the demonstration environment. In production, this role model should be connected to secure authentication such as institutional single sign-on, multi-factor authentication for administrators, password policies, session expiry, and server-side token validation.",
            "Auditability is an ethical and technical requirement. When a transfer is requested, accepted, rejected, assigned, progressed, or completed, the system records events that can be inspected later. When a bed changes from available to transfer assigned or occupied, the bed lifecycle table records the previous and new status, reason, patient ID, transfer ID, and actor. Audit logs also capture broader entity actions. This is important because emergency workflows can later be reviewed for quality improvement, accountability, and dispute resolution. Without event history, a system might make operations faster but less transparent. The project therefore treats traceability as part of responsible design.",
            "Ethically, the system must support human judgement rather than override it. A recommendation score should not become an unquestioned command. The implemented design helps by exposing reasons, scores, route sources, model names, and risk factors. However, future deployment should include explicit override workflows. For example, a receiving hospital might reject a technically suitable transfer because a specialist is unavailable, an ICU bay is physically blocked, or a patient requires a service not captured by the current input fields. The interface should allow staff to document such reasons. This is consistent with clinical decision-support principles: the software should inform, structure, and record decisions while preserving professional responsibility.",
            "Deployment planning begins with the local architecture already present in the project. The prototype can be started through Windows scripts and can run backend, frontend, database initialization, routing model warmup, and fleet simulation on one development machine. Docker Compose is included for a more reproducible environment. A production version would separate concerns: a frontend static hosting layer, a FastAPI application server, a PostgreSQL/PostGIS database, an OSRM or Valhalla routing service, a Redis or message broker layer for real-time events, and secure integration services for hospital systems and ambulance GPS. Logs, backups, monitoring, and health checks would be required before real operation.",
            "The database deployment path is clear. SQLite supports the local MVP because it is simple, file-based, and easy to reset during development. PostgreSQL with PostGIS is the correct production direction because it can store geographic points and route geometries, support stronger concurrency, and integrate with geospatial queries. The documentation already names the production target and provides a DATABASE_URL pattern. Migration tooling should be added in future work so schema changes can be applied safely. For a hospital environment, database backups, point-in-time recovery, encryption at rest, access control, and separation of development and production datasets would be mandatory.",
            "The routing deployment path also needs careful planning. The local GraphML approach is excellent for a final-year prototype because it avoids external service dependency and demonstrates graph algorithms clearly. At larger scale, loading and searching a NetworkX graph inside the API process may become slow or memory heavy. A production system could run OSRM, Valhalla, GraphHopper, or pgRouting as a dedicated routing engine. The API would then ask the routing engine for candidate paths and apply the project's urgency-aware risk scoring on top. This separation would improve performance while preserving the research contribution: combining traffic, risk, capacity, and urgency into transparent transfer recommendations.",
            "Project management followed incremental delivery. The earliest milestone was a research MVP with seeded hospitals, rule-based urgency, route ranking, FastAPI endpoints, and a React map. Later work added database-backed users, ambulances, ICU beds, patient packets, transfer requests, event histories, and admin workflows. The project then expanded into OSM node routing, ML traffic features, automated dispatch, live fleet simulation, capacity forecast, and scenario analytics. This sequence was appropriate because each phase created a working system before adding complexity. It also reduced risk: if ML artifacts or OSM routing were unavailable, the basic transfer workflow still worked through deterministic fallbacks.",
            "The project risks were both technical and domain-specific. Technical risks included graph routing failure, missing model artifacts, unavailable Google API keys, frontend/backend port mismatch, large training artifacts, stale ambulance locations, and inconsistent bed counts. The implementation mitigated these risks through fallback route sources, feature lookup fallbacks, startup health checks, active API base candidates, disk-space checks in traffic training, stale-location penalties in dispatch, and bed-count synchronization from bed states. Domain risks included unsafe recommendations, lack of real clinical validation, and privacy concerns. These were mitigated through hard filtering rules, explanation output, scope limitations, and thesis-level disclosure of non-production status.",
            "The research contribution can be stated at three levels. At the application level, the project delivers a complete prototype that supports the full transfer lifecycle rather than only a recommendation screen. At the algorithmic level, it demonstrates an urgency-aware multi-objective route and dispatch strategy that combines time, risk, distance, capacity, and coverage. At the software engineering level, it shows how healthcare decision support can be built as modular services with typed schemas, persistent events, role-aware workflows, and map-based visualization. The combination of these three levels makes the project suitable for a final-year Computer Science thesis because it contains both practical engineering and research-oriented reasoning.",
            "From a learning perspective, the project demonstrates mastery of several Computer Science areas. Database design appears in normalized operational tables and lifecycle histories. Algorithms appear in graph search, scoring, ranking, forecasting, and simulation. Machine learning appears in urgency and traffic training pipelines. Software architecture appears in route-service-model separation and fallback design. Human-computer interaction appears in the operational dashboard and driver map. Software testing appears in endpoint checks and workflow validation. These areas are connected through one domain problem, which is stronger than presenting them as disconnected assignments. The final system is therefore both a healthcare prototype and a portfolio of applied computing skills.",
            "Documentation was treated as an engineering artifact, not only as a final academic requirement. The repository includes a project blueprint, API contract, database schema, advanced traffic model note, testing guide, ML module documentation, startup scripts, environment examples, and this thesis. These files help another student, supervisor, or evaluator understand how to run the system, what endpoints exist, what models are expected, what fallbacks are available, and how the final prototype differs from the original MVP. Good documentation is especially important in a multi-module healthcare project because the code spans backend services, frontend workflows, data files, ML scripts, routing assets, and operational assumptions.",
            "Maintainability was supported through typed request and response schemas. Pydantic models define the shape of patient conditions, urgency prediction responses, route options, traffic status, hospital updates, ICU bed updates, ambulance summaries, transfer summaries, dashboard summaries, capacity forecasts, simulation summaries, and recommendations. This reduces accidental mismatch between frontend and backend. TypeScript types in the frontend mirror the backend data structures, making it easier to reason about route payloads, patient packets, dispatch results, simulation outputs, and mission data. In a future team environment, generated API clients or OpenAPI-based type generation could further reduce duplication.",
            "The project also demonstrates resilience as a design value. The frontend tries multiple backend base URLs, displays an offline message, and continues reconnecting. The backend warms models but does not block the whole application if optional routes fail. The routing stack can use local OSM, live Google Routes, cached trained models, or direct fallback geometry. The fleet simulator catches internal errors so it cannot crash the clinical API. Training scripts check for missing dependencies, missing columns, and insufficient disk space. These implementation choices show that the project was built with real demonstration risk in mind, not only ideal-case execution.",
            "For viva or presentation purposes, the strongest demonstration path is to show one complete transfer. The presenter can begin with the hospital network map, explain the bed and ambulance data, enter a critical patient condition, generate recommendations, compare route risk and ETA, create a transfer, accept it as the destination hospital, inspect dispatch scoring, switch to the ambulance crew, follow the route, complete the case, and then verify that the bed and event timeline updated. This walkthrough proves the integration of every major module in a few minutes. It also gives examiners concrete evidence that the thesis chapters correspond to working software.",
        ],
    ),
]


IMPLEMENTATION_CATALOGUE = [
    ("FastAPI application", "Defines the API surface, startup lifecycle, CORS policy, health endpoint, router registration, and model warmup."),
    ("Hospital module", "Stores hospital coordinates, ICU types, capacity fields, condition capability flags, and ventilator support used by filtering."),
    ("Urgency service", "Calculates a transparent score from condition type, oxygen level, blood pressure, consciousness, and ventilator requirement."),
    ("Transfer service", "Filters hospitals, compares routes, calculates capacity bonus, creates ranked recommendation objects, and assigns ranks."),
    ("Routing service", "Combines traffic prediction, OSM graph routing, Google fallback, route scoring, route explanations, and common schema output."),
    ("OSM graph routing", "Loads Colombo GraphML, finds nearest nodes, computes weighted paths, builds polylines, steps, risk features, and route factors."),
    ("Traffic model service", "Loads scikit-learn artifacts and traffic-ready features, builds time-aware feature rows, predicts congestion and duration, and falls back safely."),
    ("Admin workflow", "Creates transfers, handles acceptance/rejection, reserves beds, dispatches ambulances, records events, and exposes dashboards."),
    ("Ambulance workflow", "Returns active missions, accepts location updates, processes pickup/dropoff actions, and finalizes delivery."),
    ("Dispatch service", "Ranks available ambulances using ETA, pickup risk, base coverage, same-base preference, stale location, and off-base cost."),
    ("ICU bed management", "Tracks individual bed states, patient assignment, lifecycle events, and synchronized hospital occupied-bed counts."),
    ("Capacity forecasting", "Projects bed availability at 1, 3, 6, and 12 hour horizons using inbound transfers, releases, and cleaning recovery."),
    ("Simulation analytics", "Runs baseline, evening surge, mass casualty, and respiratory wave scenarios with bed and ambulance pressure outputs."),
    ("Fleet simulation", "Moves available off-base ambulances back to their bases using OSM polylines while keeping them assignable."),
    ("React frontend", "Provides role-aware transfer planning, dashboards, maps, bed editing, transfer detail, event timeline, and mission control."),
    ("Leaflet map", "Displays hospital markers, ambulance markers, route polylines, popups, legends, and map fitting for administrator views."),
    ("MapLibre driver map", "Displays a pitch-and-bearing navigation view with route lines, stop markers, ambulance marker, and recenter behavior."),
    ("ML training scripts", "Generate and train urgency and traffic models, compare algorithms, save artifacts, and write evaluation reports."),
]


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for level, size, color in [(1, 16, "2E74B5"), (2, 13, "2E74B5"), (3, 12, "1F4D78")]:
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(6 if level < 3 else 4)
    return doc


def add_front_matter(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("ICU Capacity-Aware Emergency Transfer Decision Support System")
    set_font(r, size=22, bold=True, color="0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A Final Year Computer Science Thesis")
    set_font(r, size=14, italic=True, color="1F4D78")

    for text in [
        "Submitted in partial fulfilment of the requirements for the degree of Bachelor of Science in Computer Science",
        "Project Domain: Healthcare Decision Support, Emergency Routing, Machine Learning, and Full-Stack Systems",
        "Prepared for: Hospital Project Evaluation",
        "Date: June 2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, size=11)

    doc.add_paragraph()
    add_heading(doc, "Declaration", 1)
    add_para(
        doc,
        "I declare that this thesis describes the design and implementation of the hospital emergency transfer decision-support project developed as part of final year Computer Science work. The document summarizes the project objectives, architecture, algorithms, implementation details, evaluation approach, limitations, and future improvements in a professional academic format.",
    )
    add_heading(doc, "Acknowledgement", 1)
    add_para(
        doc,
        "I acknowledge the academic guidance, technical resources, and project feedback that supported the development of this prototype. The project benefited from the study of healthcare workflows, graph routing, machine learning, database design, and modern web application development.",
    )
    doc.add_page_break()


def add_static_toc(doc: Document):
    add_heading(doc, "Table of Contents", 1)
    items = [
        "Abstract",
        "Chapter 1: Introduction",
        "Chapter 2: Background and Related Work",
        "Chapter 3: Requirements and Scope",
        "Chapter 4: Methodology",
        "Chapter 5: System Architecture",
        "Chapter 6: Database and Data Model",
        "Chapter 7: Algorithms and Decision Logic",
        "Chapter 8: Implementation Details",
        "Chapter 9: Testing and Evaluation",
        "Chapter 10: Results, Limitations, and Future Work",
        "Chapter 11: Security, Ethics, Deployment, and Project Management",
        "Implementation Catalogue",
        "References",
        "Appendix A: API Summary",
        "Appendix B: Demonstration Workflow",
    ]
    for item in items:
        add_para(doc, item)
    doc.add_page_break()


def add_catalogue(doc: Document):
    add_heading(doc, "Implementation Catalogue", 1)
    add_para(
        doc,
        "This catalogue records the system at micro scale. Each item identifies a developed component and the role it plays in the final prototype. The purpose is to show that the project is not a single-page demonstration but a complete software system with interacting services, data models, algorithms, and user interfaces.",
    )
    rows = [[name, detail] for name, detail in IMPLEMENTATION_CATALOGUE]
    add_table(doc, ["Component", "Implemented Responsibility"], rows, [2.0, 4.4])
    for name, detail in IMPLEMENTATION_CATALOGUE:
        add_heading(doc, name, 2)
        add_para(
            doc,
            f"The {name.lower()} was implemented to support the hospital transfer decision workflow in a concrete and testable way. {detail} In the overall architecture it contributes to one of three responsibilities: improving clinical suitability, improving operational coordination, or improving route and fleet intelligence. This component is also designed with fallback behavior or transparent output so that a user can understand what the system is doing during a transfer decision.",
        )


def add_references(doc: Document):
    add_heading(doc, "References", 1)
    refs = [
        "FastAPI Documentation. API framework concepts, routing, dependency injection, and OpenAPI generation.",
        "React Documentation. Component state, hooks, rendering, and single-page application development.",
        "SQLAlchemy Documentation. ORM mapping, sessions, relationships, and database abstraction.",
        "NetworkX Documentation. Graph structures and shortest path algorithms for weighted networks.",
        "OpenStreetMap Project. Open geospatial data used for road-network representation.",
        "scikit-learn Documentation. Machine learning pipelines, preprocessing, classifiers, regressors, and evaluation metrics.",
        "MapLibre GL Documentation. Browser-based vector map rendering and interactive navigation views.",
        "Leaflet Documentation. Web map rendering and marker/polyline interaction.",
        "HL7 FHIR Specification. Reference healthcare interoperability model for future bed and location integration.",
    ]
    for ref in refs:
        add_bullet(doc, ref)


def add_appendices(doc: Document):
    add_heading(doc, "Appendix A: API Summary", 1)
    rows = [
        ["/api/hospitals", "List and inspect hospitals with ICU capability and capacity summaries."],
        ["/api/predictions/urgency", "Predict urgency class and score from patient condition input."],
        ["/api/routes/optimize", "Compare shortest-time and ML traffic risk-aware routes."],
        ["/api/transfers/recommend", "Return ranked receiving-hospital recommendations."],
        ["/api/admin/dashboard", "Return hospitals, ambulances, and transfers for the command dashboard."],
        ["/api/admin/transfers", "Create transfer requests with patient and handover details."],
        ["/api/admin/transfers/{id}/accept", "Accept transfer, reserve bed, and trigger ambulance dispatch."],
        ["/api/ambulance/mission", "Return active mission for an ambulance crew user."],
        ["/api/traffic/status", "Report model, feature-data, and Google API availability."],
        ["/api/admin/simulation/run", "Run capacity and fleet pressure scenario simulation."],
    ]
    add_table(doc, ["Endpoint", "Purpose"], rows, [2.4, 4.0])

    add_heading(doc, "Appendix B: Demonstration Workflow", 1)
    steps = [
        "Start the full application with START_APP.cmd so that both backend and frontend are available.",
        "Log in or select a hospital administrator role from the user selector.",
        "Choose an origin hospital, required ICU type, patient condition, and ventilator requirement.",
        "Run the recommendation workflow and inspect urgency, ranked hospitals, ETA, and route risk.",
        "Request a transfer to a recommended destination hospital.",
        "Switch to the destination hospital administrator and accept the pending request.",
        "Inspect the automatically assigned ambulance and dispatch explanation.",
        "Switch to the ambulance crew user and open the mission view.",
        "Progress through pickup, patient onboard, destination travel, and completion.",
        "Return to the dashboard and verify transfer status, bed status, event timeline, and fleet status.",
    ]
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(step)
        set_font(r)


def add_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""
    run = p.add_run("ICU Transfer DSS Thesis")
    set_font(run, size=9, color="666666")
    add_page_number(p)


def count_words_docx(path: Path) -> int:
    doc = Document(path)
    texts: list[str] = []
    for p in doc.paragraphs:
        texts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return len(re.findall(r"\b[\w'-]+\b", "\n".join(texts)))


def main():
    doc = setup_document()
    add_front_matter(doc)
    add_static_toc(doc)
    for heading, paragraphs in SECTIONS:
        add_heading(doc, heading, 1)
        for paragraph in paragraphs:
            add_para(doc, paragraph)
    add_catalogue(doc)
    add_references(doc)
    add_appendices(doc)
    add_footer(doc)
    doc.save(OUT)
    print(f"saved={OUT}")
    print(f"word_count={count_words_docx(OUT)}")


if __name__ == "__main__":
    main()
