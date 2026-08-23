from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "Mac_Setup_Guide_Hospital_Project.docx"


def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run, size={1: 16, 2: 13}.get(level, 11), bold=True, color="2E74B5")


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_font(run)


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F4F7")
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(code)
    set_font(run, name="Consolas", size=10)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Mac Setup Guide for Hospital Project")
    set_font(run, size=22, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("GitHub-based development workflow for macOS and Windows")
    set_font(run, size=12, italic=True, color="1F4D78")

    add_para(
        doc,
        "This document explains how to set up the Hospital ICU Transfer DSS project on a Mac after the project has been pushed to GitHub. Follow these steps so you can develop the same project from both your Mac and Windows PC.",
    )

    add_heading(doc, "1. Install Required Tools")
    add_para(doc, "First install Apple command line tools:")
    add_code(doc, "xcode-select --install")
    add_para(doc, "Install Homebrew if it is not already installed:")
    add_code(doc, '/bin/bash -c "$(curl -fsSL https://raw.githubusercontent.com/Homebrew/install/HEAD/install.sh)"')
    add_para(doc, "Then install Git, Python, and Node.js:")
    add_code(doc, "brew install git python node")
    add_para(doc, "Check that the tools are installed correctly:")
    add_code(doc, "git --version\npython3 --version\nnode --version\nnpm --version")

    add_heading(doc, "2. Clone the GitHub Project")
    add_para(doc, "Use a normal development folder. Do not place the project inside iCloud Drive because cloud syncing can conflict with Git, Python virtual environments, and Node dependencies.")
    add_code(doc, "mkdir -p ~/Projects\ncd ~/Projects\ngit clone https://github.com/nisalanthony/Hospital.git\ncd Hospital")

    add_heading(doc, "3. Set Up the Backend")
    add_para(doc, "Create and activate a Python virtual environment inside the backend folder:")
    add_code(doc, "cd ~/Projects/Hospital/backend\npython3 -m venv .venv\nsource .venv/bin/activate")
    add_para(doc, "Install backend packages:")
    add_code(doc, "pip install --upgrade pip\npip install -r requirements.txt")
    add_para(doc, "Start the FastAPI backend on port 8001:")
    add_code(doc, "uvicorn app.main:app --reload --port 8001")
    add_para(doc, "Keep this terminal open. The backend should be available at:")
    add_code(doc, "http://127.0.0.1:8001")
    add_para(doc, "You can test the API documentation page in your browser:")
    add_code(doc, "http://127.0.0.1:8001/docs")

    add_heading(doc, "4. Set Up the Frontend")
    add_para(doc, "Open a second terminal and install the frontend dependencies:")
    add_code(doc, "cd ~/Projects/Hospital/frontend\nnpm install")
    add_para(doc, "Start the React/Vite frontend:")
    add_code(doc, "npm run dev")
    add_para(doc, "Open the frontend in your browser:")
    add_code(doc, "http://127.0.0.1:5173")

    add_heading(doc, "5. Normal Daily Git Workflow")
    add_para(doc, "Before you start work on either device, pull the latest changes:")
    add_code(doc, "cd ~/Projects/Hospital\ngit pull")
    add_para(doc, "After making changes, commit and push them:")
    add_code(doc, 'git status\ngit add .\ngit commit -m "Describe your change"\ngit push')
    add_para(doc, "Then on the other device, run:")
    add_code(doc, "git pull")

    add_heading(doc, "6. Important Notes")
    for item in [
        "Do not keep the project inside iCloud Drive, OneDrive, Google Drive, or Dropbox.",
        "Use ~/Projects/Hospital on the Mac for the local clone.",
        "Use GitHub as the source of truth between Windows and Mac.",
        "Do not commit .venv, node_modules, .env, database files, logs, or cache files.",
        "The project .gitignore already protects common local-only files.",
        "If you work on both devices at the same time, commit and push from one device before pulling on the other, or use separate branches.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. Optional Mac Startup Script")
    add_para(doc, "After backend and frontend setup works, you can create a Mac startup script such as START_APP_MAC.sh to start both services more easily. The manual two-terminal method above is recommended first so that installation errors are easy to see.")
    add_code(doc, "./START_APP_MAC.sh")

    add_heading(doc, "Quick Checklist")
    for item in [
        "Install Xcode command line tools.",
        "Install Homebrew.",
        "Install Git, Python, and Node.",
        "Clone the GitHub repository.",
        "Create backend virtual environment.",
        "Install backend requirements.",
        "Run FastAPI backend on port 8001.",
        "Install frontend npm packages.",
        "Run Vite frontend on port 5173.",
        "Use git pull, commit, and push to sync between Mac and Windows.",
    ]:
        add_bullet(doc, item)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
